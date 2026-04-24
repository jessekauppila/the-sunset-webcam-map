"""Shared python-docx styling helpers for art-grant-skill scripts.

These are the low-level python-docx utilities that both grant-application
builders and CV builders need. Higher-level content assembly (section
ordering, specific tables) lives in each script.

Import these in a builder script like:

    from _style import set_cell_border, shade_cell, configure_document
"""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor


# ---------------------------------------------------------------------------
# Table cell styling (direct XML — python-docx doesn't expose these directly)
# ---------------------------------------------------------------------------


def set_cell_border(cell, *, top=False, bottom=False, left=False, right=False,
                    color: str = "BFBFBF", size: str = "4") -> None:
    """Apply single-line borders to the specified edges of a table cell.

    Color is hex without #. Size is in eighths of a point (4 = 0.5pt).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    edges = {"top": top, "left": left, "bottom": bottom, "right": right}
    for edge, enabled in edges.items():
        if not enabled:
            continue
        tag = qn(f"w:{edge}")
        element = tcBorders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def shade_cell(cell, color_hex: str) -> None:
    """Fill a table cell with a solid color. Hex without #."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Paragraph-level styling
# ---------------------------------------------------------------------------


def add_horizontal_rule(paragraph) -> None:
    """Attach a thin gray bottom-border to a paragraph.

    Useful for section-heading rules in CVs and other horizontal separators.
    """
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "555555")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ---------------------------------------------------------------------------
# Document baselines
# ---------------------------------------------------------------------------


def configure_document(
    doc: Document,
    *,
    font_name: str = "Calibri",
    font_size_pt: int = 11,
    line_spacing: float = 1.15,
    space_after_pt: int = 6,
    margin_inches: float = 1.0,
) -> None:
    """Set page size, margins, and the default paragraph style.

    Defaults match the LACMA grant-application look (Calibri 11, 1" margins,
    1.15 line spacing). Override for CVs or other document types.
    """
    for section in doc.sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    pf = style.paragraph_format
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = line_spacing


def configure_headings(
    doc: Document,
    *,
    font_name: str = "Calibri",
    sizes: dict[str, int] | None = None,
    color: RGBColor = RGBColor(0x1F, 0x1F, 0x1F),
) -> None:
    """Configure heading styles (H1/H2/H3). sizes keyed by style name."""
    defaults = {"Title": 24, "Heading 1": 16, "Heading 2": 13, "Heading 3": 11}
    sizes = {**defaults, **(sizes or {})}
    for name, size in sizes.items():
        if name in doc.styles:
            s = doc.styles[name]
            s.font.name = font_name
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = color


# ---------------------------------------------------------------------------
# Quick content helpers
# ---------------------------------------------------------------------------


def add_heading(doc: Document, text: str, level: int = 1, font_name: str = "Calibri") -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = font_name


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    """Add a plain paragraph. One call = one visible paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_para_with_lead(doc: Document, lead: str, rest: str) -> None:
    """Add a paragraph where the opening phrase is bold. Useful for 'budget notes' style blocks."""
    p = doc.add_paragraph()
    r1 = p.add_run(lead)
    r1.bold = True
    p.add_run(rest)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
