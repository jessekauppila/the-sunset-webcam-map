"""Starter template for a grant-application .docx builder.

Copy this file, rename it to `build_<grant-name>_docx.py`, and fill in the
content dicts. See `build_lacma_docx_v2.py` in this same folder for a
complete worked example.

Run with:

    python3 build_<grant-name>_docx.py

The script writes a single .docx to the OUTPUT path.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt

# Import the shared helpers. If you move this script outside art-grant-skill/scripts/,
# adjust the import path or copy _style.py next to it.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from _style import (  # noqa: E402
    add_bullets,
    add_heading,
    add_para,
    add_para_with_lead,
    configure_document,
    configure_headings,
    set_cell_border,
    shade_cell,
)


# ---------------------------------------------------------------------------
# CONFIGURE FOR YOUR GRANT
# ---------------------------------------------------------------------------

GRANT_NAME = "Your Grant Name Here"
APPLICANT = "Your Name"
DEADLINE = "Month Day, Year · 11:59 PM Pacific"
MAX_REQUEST = "$50,000"

OUTPUT = (
    Path(__file__).resolve().parent.parent.parent
    / "Supporting Text"
    / "output"
    / "YourGrant_Application.docx"
)


# ---------------------------------------------------------------------------
# BUDGET ROWS — 3-tuple (category, item, cost) with is_total flag
# ---------------------------------------------------------------------------

BUDGET_ROWS: list[tuple[str, str, str, bool]] = [
    # (category, item, cost_string, is_total_row)
    # Example:
    # ("Principal Fees", "Your name — 24 months", "$12,000", False),
    # ("TOTAL", "", "$50,000", True),
]


# ---------------------------------------------------------------------------
# IMPLEMENTATION PLAN — 4-tuple (phase, timeline, milestones, cost)
# ---------------------------------------------------------------------------

PLAN_ROWS: list[tuple[str, str, str, str]] = [
    # (phase, timeline, milestones_description, cost_string)
    # Example:
    # ("1. Foundation", "Fall 2026", "Kick off ML architecture; ...", "$11,000"),
]


# ---------------------------------------------------------------------------
# TABLE BUILDERS — standard shapes reused across grant applications
# ---------------------------------------------------------------------------


def add_budget_table(doc: Document, rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(1.8), Inches(4.2), Inches(0.8)]

    hdr_cells = table.rows[0].cells
    for i, (cell, header) in enumerate(zip(hdr_cells, ["Category", "Item", "Cost"])):
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        shade_cell(cell, "E8E8E8")
        cell.width = widths[i]
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)

    last_cat = None
    for cat, item, cost, is_total in rows:
        row = table.add_row().cells
        row[0].text = "" if cat == last_cat else cat
        row[1].text = item
        row[2].text = cost
        for i, cell in enumerate(row):
            cell.width = widths[i]
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
        if is_total:
            for cell in row:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.bold = True
                shade_cell(cell, "F3F3F3")
        last_cat = cat


def add_plan_table(doc: Document, rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [Inches(1.4), Inches(1.1), Inches(3.5), Inches(0.8)]

    hdr = table.rows[0].cells
    for i, (cell, header) in enumerate(zip(hdr, ["Phase", "Timeline", "Milestones", "Cost"])):
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        shade_cell(cell, "E8E8E8")
        cell.width = widths[i]
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)

    for phase, timeline, milestones, cost in rows:
        row = table.add_row().cells
        row[0].text = phase
        row[1].text = timeline
        row[2].text = milestones
        row[3].text = cost
        for i, cell in enumerate(row):
            cell.width = widths[i]
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# DOCUMENT ASSEMBLY — fill in each section per your grant's structure
# ---------------------------------------------------------------------------


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    configure_headings(doc)

    # Title block
    add_heading(doc, GRANT_NAME, level=0)
    add_para(doc, f"Applicant: {APPLICANT}", italic=True)
    add_para(doc, f"Deadline: {DEADLINE}", italic=True)
    add_para(doc, f"Amount Requested: {MAX_REQUEST}", italic=True)

    # §1 — Name of Project
    add_heading(doc, "1. Name of Project", level=1)
    add_para(doc, "Your project name here.")

    # §2 — Three Words
    add_heading(doc, "2. Three Words", level=1)
    add_para(doc, "Word. Word. Word.")

    # §3 — One-Sentence Description
    add_heading(doc, "3. One-Sentence Description", level=1)
    add_para(doc, "Your one-sentence description here.")

    # §4 — Full Description (respect the word cap)
    add_heading(doc, "4. Full Description", level=1)
    add_para(doc, "Paragraph 1 of the full description.")
    add_para(doc, "Paragraph 2 of the full description.")
    # ... add more paragraphs as needed

    # §N — Budget
    add_heading(doc, "N. Detailed Project Budget", level=1)
    if BUDGET_ROWS:
        add_budget_table(doc, BUDGET_ROWS)
    else:
        add_para(doc, "(Fill in BUDGET_ROWS at the top of this file.)", italic=True)

    # §N+1 — Implementation Plan
    add_heading(doc, "N+1. Implementation Plan", level=1)
    if PLAN_ROWS:
        add_plan_table(doc, PLAN_ROWS)
    else:
        add_para(doc, "(Fill in PLAN_ROWS at the top of this file.)", italic=True)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
