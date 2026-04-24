"""Build v5 Part 2 of the LACMA Art + Technology Lab 2026 grant application —
§11 Detailed Project Budget and §12 Implementation Plan only — as a clean Word
(.docx). The user has hand-edited §1–§10 separately; this file produces the
bottom half so the two can be combined.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent.parent  # art-grant-skill/scripts/ → repo root
OUTPUT = REPO_ROOT / "Supporting Text" / "output" / (
    "LACMA_Art_Tech_Lab_2026_Kauppila_Harris_v5_part2.docx"
)


# ---------------------------------------------------------------------------
# Style helpers (mirrored from build_lacma_docx_v2.py)
# ---------------------------------------------------------------------------


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = qn(f"w:{edge}")
            element = tcBorders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:color"), "BFBFBF")


def shade_cell(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for name, size, bold, color in [
        ("Title", 24, True, RGBColor(0x1F, 0x1F, 0x1F)),
        ("Heading 1", 16, True, RGBColor(0x1F, 0x1F, 0x1F)),
        ("Heading 2", 13, True, RGBColor(0x1F, 0x1F, 0x1F)),
        ("Heading 3", 11, True, RGBColor(0x40, 0x40, 0x40)),
    ]:
        if name in doc.styles:
            s = doc.styles[name]
            s.font.name = "Calibri"
            s.font.size = Pt(size)
            s.font.bold = bold
            s.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_para_with_lead(doc: Document, lead: str, rest: str) -> None:
    """Add a paragraph where the opening phrase is bold (used for budget notes)."""
    p = doc.add_paragraph()
    r1 = p.add_run(lead)
    r1.bold = True
    p.add_run(rest)


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------


def add_budget_table(doc: Document, rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(1.8), Inches(4.2), Inches(0.8)]
    hdr_cells = table.rows[0].cells
    headers = ["Category", "Item", "Cost"]
    for i, (cell, header) in enumerate(zip(hdr_cells, headers)):
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        shade_cell(cell, "E8E8E8")
        cell.width = widths[i]
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)

    last_cat = None
    for cat, item, cost, is_total in rows:
        row = table.add_row().cells
        row[0].text = "" if cat == last_cat else cat
        row[1].text = item
        row[2].text = cost
        for i, cell in enumerate(row):
            cell.width = widths[i]
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
        if is_total:
            for cell in row:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.bold = True
                shade_cell(cell, "F3F3F3")
        last_cat = cat


def add_plan_table(doc: Document, rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [Inches(1.4), Inches(1.1), Inches(3.5), Inches(0.8)]
    headers = ["Phase", "Timeline", "Milestones", "Cost"]
    hdr = table.rows[0].cells
    for i, (cell, text) in enumerate(zip(hdr, headers)):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        shade_cell(cell, "E8E8E8")
        cell.width = widths[i]
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)

    for phase, timeline, milestones, cost in rows:
        row = table.add_row().cells
        row[0].text = phase
        row[1].text = timeline
        row[2].text = milestones
        row[3].text = cost
        for i, cell in enumerate(row):
            cell.width = widths[i]
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------


BUDGET_ROWS = [
    # (Category, Item, Cost, is_total)
    ("Principal Fees", "Jesse Kauppila — artist + project lead (24 months)", "$12,000", False),
    ("Co-applicant Fees",
     "Kameron Decker Harris — ML co-lead (24 months): ML architecture, training strategy, model evaluation — at ⅔ of the principal fee",
     "$8,000", False),
    ("Webcam Prototyping (NA pilot)", "30 × Pi Zero 2 W", "$750", False),
    ("Webcam Prototyping (NA pilot)", "30 × Camera modules", "$600", False),
    ("Webcam Prototyping (NA pilot)", "30 × MicroSD, high-endurance", "$360", False),
    ("Webcam Prototyping (NA pilot)", "30 × Power supplies + cables", "$300", False),
    ("Webcam Prototyping (NA pilot)", "30 × Weatherproof enclosures", "$450", False),
    ("Webcam Prototyping (NA pilot)", "Custom PCB design + small run (Pi + ESP32 board variants, 2 iterations)", "$2,500", False),
    ("Webcam Prototyping (NA pilot)", "Enclosure design, 3D printing", "$1,000", False),
    ("Webcam Prototyping (NA pilot)", "Cellular / LTE modems (5–10 units)", "$500", False),
    ("Webcam Prototyping (NA pilot)", "ESP32 experimental units + peripherals", "$500", False),
    ("Webcam Prototyping (NA pilot)", "Shipping, tools, spare parts", "$540", False),
    ("Webcam Prototyping (NA pilot)", "— subtotal —", "$7,500", False),
    ("Travel (Bellingham → LA, 2 people)", "2027 Symposium trip × 2 people", "$4,400", False),
    ("Travel (Bellingham → LA, 2 people)", "2028 Demo Day trip × 2 people", "$6,800", False),
    ("Travel (Bellingham → LA, 2 people)", "Canopy Art & Iron (local drive, Bow WA)", "$300", False),
    ("Travel (Bellingham → LA, 2 people)", "— subtotal —", "$11,500", False),
    ("Software & Services",
     "Anthropic Claude — Jesse, Max tier ($200/mo × 24, for algorithm-training + Claude-Code pipeline work)",
     "$4,800", False),
    ("Software & Services", "Anthropic Claude — Kameron, Pro tier ($100/mo × 24)", "$2,400", False),
    ("Software & Services", "Anthropic + OpenAI + Google Gemini vision API calls (LLM labeling pipeline)", "$1,000", False),
    ("Software & Services", "Cursor Pro — Jesse ($20/mo × 24)", "$480", False),
    ("Software & Services", "Cursor Pro — Kameron ($20/mo × 24)", "$480", False),
    ("Software & Services", "Vercel Pro — project deployment ($20/mo × 24)", "$480", False),
    ("Software & Services", "Neon Postgres — project database (24 months)", "$600", False),
    ("Software & Services", "Mapbox GL JS (free tier covers expected usage; overage from contingency)", "$0", False),
    ("Software & Services", "Domain + SSL (sunrisesunset.studio, 2 yrs)", "$100", False),
    ("Software & Services", "— subtotal —", "$10,340", False),
    ("Contingency", "Replacement hardware, API overages, shipping", "$660", False),
    ("TOTAL", "", "$50,000", True),
]


BUDGET_NOTES = [
    ("Principal fees structured at a 3:2 ratio. ",
     "Jesse at $12,000 and Kameron at $8,000 (exactly ⅔ of the principal fee) reflect the division of labor: "
     "Jesse leads the project, builds the cameras, the web app, and the installation; Kameron leads the ML "
     "architecture, training strategy, and model evaluation. The fees are modest — roughly $500 and $333 per "
     "month respectively across 24 months — because we've deliberately shifted budget toward the research "
     "instrument (the webcam network) and toward covering both people's participation across the grant's "
     "subscription and travel obligations."),
    ("Webcam Prototyping ($7,500) funds a dual track. ",
     "The bulk (~30 Raspberry Pi Zero 2 W units, the main deployment platform) is complemented by a smaller "
     "ESP32-based experimental run. ESP32 units cost roughly 1/10 of a Pi, which lets us test whether the "
     "network can extend cheaply into harder-to-reach locations — particularly where solar power and cellular "
     "are the only options. Custom PCB design ($2,500 for two iterations, one Pi carrier and one ESP32 carrier) "
     "and enclosure work ($1,000) produce the hardware artifact that could later be offered as a Kickstarter "
     "reward. The webcam line is sized close to Kameron's fee because the network and the algorithm are the "
     "two primary research instruments the grant funds."),
    ("Software & Services ($10,340) covers both people. ",
     "Per-user subscriptions (Anthropic Claude, Cursor Pro) are budgeted for both Jesse and Kameron. Shared "
     "project infrastructure (Vercel, Neon Postgres, Mapbox, domain) is single-license. Jesse's Anthropic "
     "Claude subscription is budgeted at the $200/mo Max tier to support algorithm-training and Claude-Code-"
     "assisted pipeline development; Kameron's at $100/mo. The vision-API line ($1,000) covers runtime "
     "labeling calls to Anthropic Claude, OpenAI, and Google Gemini as part of the LLM-driven data-labeling "
     "pipeline — separate from the subscription costs and drawn on programmatically per training batch."),
    ("Travel ($11,500) covers both Jesse and Kameron. ",
     "Two round trips each to the 2027 LACMA Biennial Symposium and the 2028 Demo Day, plus a local drive to "
     "the Canopy Art & Iron rendezvous in Bow, WA. Having both the principal artist and the ML co-lead "
     "present at the Symposium and Demo Day matters: the project's public-facing story is the collaboration "
     "between artistic and computational practice, and the Q&A moments at these events are where that "
     "collaboration is most legible."),
    ("What's NOT in the budget. ",
     "We don't buy monitors, stands, wiring, or kiosk machines — venues source their own displays for formal "
     "installations, and informal installations use found screens and participant phones. We also don't "
     "budget a GPU workstation — training runs that need a GPU will use cloud credits drawn from the vision-"
     "API / Anthropic subscription lines. No documentation line: video and photo documentation will be "
     "captured alongside the existing project work rather than as a separate paid deliverable. The grant "
     "funds the research instrument (the webcam network and the algorithm), the people doing the research "
     "(Jesse and Kameron), and the subscriptions that enable their work — not the display surface, not "
     "dedicated compute, not production crews. This cost discipline aligns with the project's \"real, not "
     "generated\" ethic: the display is as found as the sunsets."),
]


PLAN_ROWS = [
    ("1. Foundation + Webcam v1",
     "Fall 2026 (Months 1–4)",
     "Formalize co-lead agreement with Kameron, kick off ML architecture; upgrade ML pipeline to image-based "
     "inference; expand training dataset; validate against human ratings (Pearson > 0.80 gate); begin outreach "
     "to Windy / EarthCam; build first 5–10 Pi Zero 2 W cameras and first ESP32 prototypes; identify NA "
     "deployment partners.",
     "$11,000"),
    ("2. Modular Install Development + NA Webcam Rollout",
     "Winter 2026–27 (Months 4–8)",
     "Refine the modular installation approach (venue-sourced monitor mix + participant phone-ring protocol); "
     "develop the public rating app; iterate webcam hardware (custom PCB, weatherproof enclosure, remote-"
     "tuning firmware, ESP32 field tests); deploy 8–10 cameras across North American locations; exhibit the "
     "two-screen diptych at Canopy Art & Iron's annual rendezvous (Bow, WA).",
     "$13,000"),
    ("3. LACMA Symposium Demo",
     "Spring 2027 (Months 8–12)",
     "Install and demonstrate at the 2027 Biennial Symposium (venue-sourced monitors + phone-ring demo); "
     "present mid-term findings (model evolution, human–machine rating divergence, pilot-webcam data, "
     "Canopy exhibit reflections); travel Bellingham → LA (2 people).",
     "$9,000"),
    ("4. Refinement + Kickstarter Prep + Global Network",
     "Summer–Fall 2027 (Months 12–18)",
     "Incorporate symposium feedback; train second-generation model on expanded dataset (gallery ratings + "
     "rating-app data + pilot imagery); finalize Kickstarter-ready camera and campaign; distribute cameras "
     "internationally; open-source release of ML pipeline.",
     "$8,000"),
    ("5. Final Installation + 2028 Demo Day",
     "Winter–Spring 2028 (Months 18–24)",
     "Full-scale modular installation for 2028 Demo Day (venue-sourced monitor ring + phone-ring public "
     "event); published dataset, model weights, camera hardware plans; travel Bellingham → LA (2 people).",
     "$9,000"),
]


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    # §11
    add_heading(doc, "11. Detailed Project Budget", level=1)
    add_budget_table(doc, BUDGET_ROWS)

    doc.add_paragraph()  # spacing

    add_heading(doc, "Budget notes", level=2)
    for lead, rest in BUDGET_NOTES:
        add_para_with_lead(doc, lead, rest)

    # §12
    doc.add_paragraph()
    add_heading(doc, "12. Implementation Plan", level=1)
    add_plan_table(doc, PLAN_ROWS)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
