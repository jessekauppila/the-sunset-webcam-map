"""Starter template for an artist/exhibition CV .docx builder.

Copy this file, rename it to `build_<name>_cv.py`, and fill in the
content dicts. See `build_cv_docx.py` in this same folder for a worked
example (Jesse Kauppila's CV).

Run with:

    python3 build_<name>_cv.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Import shared helpers from alongside this file.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from _style import (  # noqa: E402
    add_horizontal_rule,
    configure_document,
)


# ---------------------------------------------------------------------------
# CONFIGURE FOR YOUR CV
# ---------------------------------------------------------------------------

NAME = "YOUR NAME"
CONTACT_LINE = "yourwebsite.com  ·  github.com/yourusername"

BASE_FONT = "Garamond"
BODY_SIZE = Pt(11)

OUTPUT = (
    Path(__file__).resolve().parent.parent.parent
    / "Supporting Text"
    / "output"
    / "YourName_CV.docx"
)


# ---------------------------------------------------------------------------
# CV CONTENT — each section is a list of entries
# ---------------------------------------------------------------------------

EDUCATION: list[tuple[str, str, str]] = [
    # (year_range, degree, institution)
    # ("2013–2016", "MFA", "Carnegie Mellon University"),
    # ("2004–2007", "BA", "Reed College"),
]

PROFESSIONAL_EXPERIENCE: list[tuple[str, str, str]] = [
    # (year_range, title, organization + short description)
    # ("2025–present", "Full-Stack Developer", "Vetta AI — customer-facing UI for AI talent-matching"),
]

SELECTED_WORK: list[tuple[str, str]] = [
    # (year, description)
    # ("2024", "Interactive water exhibit for Pittsburgh Children's Museum"),
]

GRANTS_AWARDS: list[tuple[str, str]] = [
    # (year, award_name)
]

VOLUNTEER: list[tuple[str, str]] = [
    # (year_range, role + organization)
]


# ---------------------------------------------------------------------------
# LAYOUT HELPERS
# ---------------------------------------------------------------------------


def add_contact_header(doc: Document) -> None:
    """Large centered name + italic contact line."""
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(NAME)
    name_run.font.name = BASE_FONT
    name_run.font.size = Pt(22)
    name_run.font.bold = True

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(18)
    contact_run = contact_p.add_run(CONTACT_LINE)
    contact_run.font.name = BASE_FONT
    contact_run.font.size = Pt(10)
    contact_run.font.italic = True


def add_section_heading(doc: Document, text: str) -> None:
    """Small-caps uppercase heading with a thin rule beneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.name = BASE_FONT
    run.font.size = Pt(11)
    run.font.bold = True
    add_horizontal_rule(p)


def add_year_entry(doc: Document, year: str, line: str) -> None:
    """A single entry — year on the left, description following."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    year_run = p.add_run(f"{year}    ")
    year_run.font.name = BASE_FONT
    year_run.font.size = BODY_SIZE
    year_run.font.bold = True

    desc_run = p.add_run(line)
    desc_run.font.name = BASE_FONT
    desc_run.font.size = BODY_SIZE


# ---------------------------------------------------------------------------
# DOCUMENT ASSEMBLY
# ---------------------------------------------------------------------------


def build_document() -> Document:
    doc = Document()
    configure_document(doc, font_name=BASE_FONT, font_size_pt=11, space_after_pt=3)

    add_contact_header(doc)

    if EDUCATION:
        add_section_heading(doc, "Education")
        for year, degree, institution in EDUCATION:
            add_year_entry(doc, year, f"{degree}, {institution}")

    if PROFESSIONAL_EXPERIENCE:
        add_section_heading(doc, "Professional Experience")
        for year, title, org in PROFESSIONAL_EXPERIENCE:
            add_year_entry(doc, year, f"{title} — {org}")

    if SELECTED_WORK:
        add_section_heading(doc, "Selected Work")
        for year, desc in SELECTED_WORK:
            add_year_entry(doc, year, desc)

    if GRANTS_AWARDS:
        add_section_heading(doc, "Grants & Awards")
        for year, award in GRANTS_AWARDS:
            add_year_entry(doc, year, award)

    if VOLUNTEER:
        add_section_heading(doc, "Volunteer")
        for year, role in VOLUNTEER:
            add_year_entry(doc, year, role)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
