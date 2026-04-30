"""Build a formatted artist-exhibition CV as a .docx.

Modeled on scripts/build_lacma_docx_v2.py. Run:

    python3 scripts/build_cv_docx.py

Writes to Supporting Text/output/Kauppila_CV_formatted.docx.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches


REPO_ROOT = Path(__file__).resolve().parent.parent.parent  # art-grant-skill/scripts/ → repo root
OUTPUT = REPO_ROOT / "Supporting Text" / "output" / "Kauppila_CV_formatted.docx"

BASE_FONT = "Garamond"
BODY_SIZE = Pt(11)
YEAR_TAB = Inches(0.65)


def configure_document(doc: Document) -> None:
    """Set page margins and the default paragraph style."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(3)


def _add_horizontal_rule(paragraph) -> None:
    """Attach a bottom border to the paragraph — a thin gray rule."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")  # 0.5 pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "555555")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_contact_header(doc: Document) -> None:
    """Render the name and contact line at the top of the document."""
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run("JESSE KAUPPILA")
    name_run.font.name = BASE_FONT
    name_run.font.size = Pt(22)
    name_run.font.bold = True

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(18)
    contact_run = contact_p.add_run(
        "jessekauppila.art  ·  github.com/jessekauppila"
    )
    contact_run.font.name = BASE_FONT
    contact_run.font.size = Pt(10)
    contact_run.font.italic = True


def add_section_heading(doc: Document, text: str) -> None:
    """Small-caps heading with a thin gray rule beneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = BASE_FONT
    run.font.size = Pt(12)
    run.font.bold = True
    _add_horizontal_rule(p)


def _set_year_tab(paragraph) -> None:
    """Attach a single left tab stop at YEAR_TAB."""
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        YEAR_TAB, alignment=WD_TAB_ALIGNMENT.LEFT
    )
    paragraph.paragraph_format.left_indent = Inches(0)


def add_entry(
    doc: Document,
    year: str,
    title: str,
    extra_lines: list[str] | None = None,
    *,
    italic_title: bool = True,
) -> None:
    """Render one entry with year in the gutter and an italic title.

    If `year` is empty, the gutter is left blank but alignment is preserved
    (classic CV convention for stacking multiple entries under one year).

    `extra_lines` become continuation paragraphs with a hanging indent at
    YEAR_TAB.
    """
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(0)
    _set_year_tab(title_p)
    if year:
        year_run = title_p.add_run(f"{year}\t")
        year_run.font.name = BASE_FONT
        year_run.font.size = BODY_SIZE
    else:
        title_p.add_run("\t")
    title_run = title_p.add_run(title)
    title_run.font.name = BASE_FONT
    title_run.font.size = BODY_SIZE
    title_run.font.italic = bool(italic_title)

    for line in extra_lines or []:
        cont = doc.add_paragraph()
        cont.paragraph_format.left_indent = YEAR_TAB
        cont.paragraph_format.space_after = Pt(0)
        cont_run = cont.add_run(line)
        cont_run.font.name = BASE_FONT
        cont_run.font.size = BODY_SIZE

    # Pt(6) empty paragraph acts as a small inter-entry spacer.
    trailer = doc.add_paragraph()
    trailer.paragraph_format.space_after = Pt(0)
    trailer.paragraph_format.space_before = Pt(0)
    trailer_run = trailer.add_run("")
    trailer_run.font.size = Pt(6)


def add_plain_entry(doc: Document, lines: list[str]) -> None:
    """Multi-line entry with no title and no year (press / publications).

    First line gets no indent; subsequent lines are indented to YEAR_TAB
    so the block reads as a unit.
    """
    if not lines:
        return
    head = doc.add_paragraph()
    head.paragraph_format.space_after = Pt(0)
    head_run = head.add_run(lines[0])
    head_run.font.name = BASE_FONT
    head_run.font.size = BODY_SIZE

    for line in lines[1:]:
        cont = doc.add_paragraph()
        cont.paragraph_format.left_indent = YEAR_TAB
        cont.paragraph_format.space_after = Pt(0)
        cont_run = cont.add_run(line)
        cont_run.font.name = BASE_FONT
        cont_run.font.size = BODY_SIZE

    # Pt(6) empty paragraph acts as a small inter-entry spacer.
    trailer = doc.add_paragraph()
    trailer.paragraph_format.space_after = Pt(0)
    trailer.paragraph_format.space_before = Pt(0)
    trailer_run = trailer.add_run("")
    trailer_run.font.size = Pt(6)


EDUCATION: list[tuple[str, list[str]]] = [
    ("Carnegie Mellon University (MFA, 2016)", ["Pittsburgh, PA"]),
    ("Reed College (BA)", ["Portland, OR"]),
    ("Hampshire College", ["Amherst, MA"]),
    ("Il Bisonte: Foundation for the Study of Printmaking",
     ["Florence, Italy"]),
    ("St. Johnsbury Academy", ["St. Johnsbury, VT"]),
    ("St. Coleman’s College", ["Co. Cork, Ireland"]),
]


# Tuples are (year, title, extra_lines). Years repeat in source order;
# where source lists two entries under one year, we pass an empty year
# on the follow-up entry so the gutter shows blank but alignment stays.
AWARDS: list[tuple[str, str, list[str]]] = [
    ("2015", "Public Art Commission: Fallow Grounds for Sculpture (2015)",
     ["Neu Kirche Contemporary Art Center"]),
    ("", "Neighbor to Neighbor Grant (2015)", ["Sprout Fund"]),
    ("", "Corrigan “Wrong Way” Travel Grant (2015)",
     ["Carnegie Mellon Department of Fine Art"]),
    # Typo fix: source missing closing paren on "(2015".
    ("", "ProSEED/Crosswalk Grant (2015)",
     ["Carnegie Mellon Fellowships and Awards"]),
    ("", "Graduate Student Research Grant (2015)",
     ["Carnegie Mellon Graduate Student Assembly"]),
    ("2014", "Tough Art Residency (2014)",
     ["The Children’s Museum of Pittsburgh"]),
    ("", "Fellowship (2014)", ["Mildred’s Lane"]),
    ("2013", "Frank-Ratchye Fund for Art @ the Frontier Microgrant (2013)",
     ["Carnegie Mellon University"]),
    ("", "Graduate Student Research Grant (2013)",
     ["Carnegie Mellon Graduate Student Assembly"]),
    ("", "Graduate Student Travel Grant (2013)",
     ["Carnegie Mellon Graduate Student Assembly"]),
    ("", "Artist in Residence (2013)", ["Rayko Photo Center"]),
    ("2012", "Artist in Residence (2009–2012)", ["Kala Art Institute"]),
    ("", "Alternative Exposure Grant for “Art for a Democratic "
     "Society” (2012)",
     ["The Andy Warhol Foundation, Southern Exposure"]),
    ("2007", "Undergraduate Research Grant (2007)", ["Reed College"]),
]


# (year, title, extra_lines). Titles quoted with curly double quotes.
SOLO_SHOWS: list[tuple[str, str, list[str]]] = [
    ("2015", "“Games”",
     ["Hyptique Pop-Up with Hannah Epstein",
      "Pittsburgh, PA"]),
    ("", "“Webs and Reticulations: Structuring Metaphors and "
     "Materials”",
     ["Posner Center, Carnegie Mellon University",
      "Curatorial project with Mary Kay Johnsen",
      "Pittsburgh, PA"]),
    ("2014", "“Screens”",
     ["Red Door Gallery, Carnegie Mellon University",
      "Pittsburgh, PA"]),
    ("2013", "“Inside Out Printer, Improvised Explosive Device "
     "(I.O.P. I.E.D.)”",
     ["Rayko Photo Center, curated by Ann Jastrab",
      "San Francisco, CA"]),
    ("2010", "“Remastering the Anthology of American Folk Music”",
     ["PLAySPACE, California College of the Arts, "
      "curated by Amanda Hunt",
      "San Francisco, CA"]),
]


GROUP_SHOWS: list[tuple[str, str, list[str]]] = [
    ("2025", "“Blanchard Mountain Rendezvous / Canopy”",
     ["Canopy Art & Iron, Bow, WA"]),
    ("", "“Fly”", ["Terramor, Bow, WA"]),

    ("2017", "“Mind Control”",
     ["AlterSpace, San Francisco, CA"]),

    ("2016", "“Corte Madera Centennial Art Exhibition”",
     ["Corte Madera Community Center, Corte Madera, CA"]),
    # Typo fix: "Miller Galler " -> "Miller Gallery" in source.
    ("", "“Self-Driving Car”",
     ["Carnegie Mellon Miller Gallery",
      "Bolinas, CA"]),
    ("", "“Process Photography / Witchy Shit”",
     ["Gospel Flats, Bolinas, CA"]),

    ("2015", "“Performance Hour”",
     ["Neu Kirche Contemporary Art Center, Pittsburgh, PA"]),
    ("", "“The Labor Portraits of Mildred’s Lane”",
     ["The Mildred Complexity, Narrowsburg, NY"]),
    ("", "“Fallow Grounds for Sculpture”",
     ["Neu Kirche Contemporary Art Center, curated by Oreen Cohen",
      "Pittsburgh, PA"]),
    ("", "“10 Minute Play Festival”",
     ["Boom Concepts, Pittsburgh, PA"]),
    ("", "“PhAb Now!”",
     ["Pittsburgh Filmmakers, curated by Adam Welch",
      "Pittsburgh, PA"]),

    ("2014", "“Humanufactory(ng) Workstyles”",
     ["School of the Art Institute of Chicago, "
      "curated by Mary Jane Jacobs and Kate Zeller",
      "Chicago, IL"]),
    ("", "“Subterraneans”",
     ["Leeds College of Art and Design, curated by Terence Jones",
      "Leeds, United Kingdom"]),
    ("", "“Tough Art”",
     ["Children’s Museum of Pittsburgh, Pittsburgh, PA"]),
    ("", "“Encountering the Unseen: Puppet Activated Lecture on "
     "the Microbiome”",
     ["Phipps Conservatory, Pittsburgh, PA"]),
    ("", "“LunarmagmaoceanLove”",
     ["NURTUREart Gallery, curated by Jaewook Lee",
      "Brooklyn, NY"]),
    ("", "“GDP”", ["The Mine Factory, Pittsburgh, PA"]),

    ("2013", "“The Making is a Re-Making”",
     ["Kala Art Institute, curated by Mayumi Hamanaka",
      "Berkeley, CA"]),

    ("2012", "“Edicola”",
     ["Colpa Press, San Francisco, CA"]),
    ("", "“Monthly Programming”",
     ["with the collective Art for a Democratic Society",
      "Art Murmur, Oakland, CA"]),
    ("", "“In Formation”",
     ["Berkeley Central Arts, curated by Amanda Curreri",
      "Berkeley, CA"]),

    ("2011", "“Books, Prints, and Things”",
     ["Wire + Nail, San Francisco, CA"]),
    ("", "“Artist Annual”",
     ["Kala Art Institute, Berkeley, CA"]),
    ("", "“Art Science Fair”",
     ["The Lab, San Francisco, CA"]),
    ("", "“Cashing Out”",
     ["Kala Art Institute, curated by Julio Cesar Morales",
      "Berkeley, CA"]),
    ("", "“Proof”",
     ["Southern Exposure, San Francisco, CA"]),
    ("", "“Experimental Notation”",
     ["MacArthur b Arthur, Oakland, CA"]),
    ("", "“Night Market”",
     ["MassMOCA, curated by James Voorhies",
      "North Adams, MA"]),
    ("", "“Artcards Presents: Performance”",
     ["The Invisible Dog, curated by Helen Homan Wu",
      "Brooklyn, NY"]),
    ("", "“Cries of San Francisco”",
     ["Southern Exposure, curated by Allison Smith and Courtney Fink",
      "San Francisco, CA"]),
    ("", "“Moonlight, Mai Tais, and Magic”",
     ["Allegra LaViola Gallery, New York, NY"]),
    ("", "“Vermont Printmakers”",
     ["Gato Nero Gallery, St. Johnsbury, VT"]),

    ("2010", "“Fresh Work”",
     ["Kala Art Institute, Berkeley, CA"]),
    ("", "“Sights + Sounds”",
     ["Noma Gallery, San Francisco, CA"]),
    ("", "“The Wassaic Festival”",
     ["The Wassaic Project, Wassaic, NY"]),
    ("", "“New Music Series”",
     ["The Luggage Store, San Francisco, CA"]),

    ("2009", "“The Living Archive”",
     ["Swell Gallery, SFAI, San Francisco, CA"]),
    ("", "“Vermont Printmakers”",
     ["Gato Nero Gallery, St. Johnsbury, VT"]),

    ("2008", "“Annual Exhibition”",
     ["Studio for Color Etching, Barga, Italy"]),
    ("", "“Aquatint Explosions”",
     ["Alt.Space Presents, Malmo, Sweden"]),
    ("", "“Bunker: Towards a Free School in the New Dark Age”",
     ["collaboration with the alt.Space Network of Artist Research "
      "Groups",
      "Hats Plus Gallery, London, UK"]),

    ("2007", "“Annual Exhibition”",
     ["Studio for Color Etching, Barga, Italy"]),
    ("", "“Learning is Fun and Dangerous”",
     ["collaboration with Red76",
      "Reed College, Portland, OR"]),

    ("2006", "“The Second Annual”",
     ["Bonnie Kahn Gallery, Portland, OR"]),
    ("", "“Tracing the Y Chromosome”",
     ["collaboration with Gerri Ondrizek",
      "Hoffman Gallery, Oregon College of Art and Craft, Portland, OR",
      "Sheehan Gallery, Whitman College, Walla Walla, WA"]),
    ("", "“Icons”",
     ["Saffron and Turmeric, Portland, OR"]),
    ("", "“Secular Confessional”",
     ["Reed Arts Week, Portland, OR"]),

    ("2005", "“Behind the Masks: Art, Culture, and History”",
     ["Southern Illinois University Museum, Carbondale, IL"]),

    ("2003", "“Il Bisonte Agli Uffizi: Vent’anni della Scuola "
     "Internazionale di Grafica d’Arte”",
     ["Galleria degli Uffizi, Florence, Italy"]),
    ("", "“Il Fino di Anno”",
     ["La Galleria di Il Bisonte, Florence, Italy"]),
]


# Each inner list is one press citation: article title then byline.
PRESS: list[list[str]] = [
    ["“Neu Kirche Contemporary Art Center.”  Mary Thomas",
     "Pittsburgh Post-Gazette, August 2015"],
    ["“‘PhAb Now’ at Pittsburgh Filmmakers Galleries.”  "
     "Kurt Shaw",
     "TribLive, July 1, 2015"],
    ["“Berkeley Central Arts Passage Unveils Its First Show.”  "
     "Alex Bigman",
     "The East Bay Express, January 2013"],
    ["“Jesse Boardman Kauppila: Italian Tartan.”  Emily Walsh",
     "Armfuls Blog, July 2012"],
    ["“‘Cries of San Francisco’: Marketplace as Art.”  "
     "Nirmala Nataraj",
     "The San Francisco Chronicle, July 2011"],
    ["“Artists Transform Downtown San Francisco into Conceptual "
     "Marketplace.”  Andy Wright",
     "The Bay Citizen"],
    ["“Jesse Kauppila at Little Paper Planes.”  R.L. Tilman",
     "Print Interesting"],
    ["“Jesse Boardman Kauppila Interview.”  Bora Mici",
     "Art Speak"],
]


PUBLICATIONS: list[list[str]] = [
    ["“Reticulations: The Evolution of the Tree Metaphor.”",
     "Pittsburgh Articulate (2015)"],
    ["“The Subterraneans.”",
     "Leeds College of Art and Design (2015)"],
    ["“Allison Smith’s Cries of San Francisco.”",
     "Southern Exposure (2013)"],
    ["“Jesse Kauppila on Remastering the Anthology of American "
     "Folk Music.”",
     "University of East Anglia, London, Conference on Harry Smith "
     "(August 2012)"],
    ["“Remastering the Anthology of American Folk Music.”",
     "Composite Arts Magazine, Issue No. 6 (January 2012)"],
    ["“Remastering, Sights + Sounds: Volume 1.”",
     "Little Paper Planes (November 2010)"],
    ["“Aquatint Explosions.”",
     "Caterwaul Quarterly (2008)"],
]


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_contact_header(doc)
    add_section_heading(doc, "Education")
    for title, extras in EDUCATION:
        add_entry(doc, year="", title=title, extra_lines=extras,
                  italic_title=False)
    # Typo fix: "Comissions" -> "Commissions" in section heading.
    add_section_heading(
        doc, "Commissions | Fellowships | Awards | Residencies"
    )
    for year, title, extras in AWARDS:
        add_entry(doc, year=year, title=title, extra_lines=extras,
                  italic_title=False)
    add_section_heading(doc, "Solo | Two-Person Shows")
    for year, title, extras in SOLO_SHOWS:
        add_entry(doc, year=year, title=title, extra_lines=extras,
                  italic_title=True)
    add_section_heading(doc, "Group Exhibitions")
    for year, title, extras in GROUP_SHOWS:
        add_entry(doc, year=year, title=title, extra_lines=extras,
                  italic_title=True)
    add_section_heading(doc, "Reviews and Press")
    for citation in PRESS:
        add_plain_entry(doc, citation)
    add_section_heading(doc, "Publications")
    for citation in PUBLICATIONS:
        add_plain_entry(doc, citation)
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
