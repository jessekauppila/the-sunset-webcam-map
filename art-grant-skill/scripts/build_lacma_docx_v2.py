"""Build v2 of the LACMA Art + Technology Lab 2026 grant application as a
clean Word (.docx).

v2 changes from build_lacma_docx.py:
- Reframed as a collaboration between Jesse Kauppila and Kameron Decker Harris
- Updated project URL to sunrisesunset.studio
- Added personal websites for both applicants
- Trimmed §4 (655 → ~498 words)
- Trimmed §8 (120 → 100 words)
- Added a Kameron Decker Harris bio paragraph
- Polished grammar in §4 / §5 / §7 / §12
- Updated §9 to remove "consulting" framing for Kameron (now co-applicant)
- Updated §11 collaborator-fees line to reflect co-lead compensation
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent.parent  # art-grant-skill/scripts/ → repo root
OUTPUT = REPO_ROOT / "Supporting Text" / "output" / (
    "LACMA_Art_Tech_Lab_2026_Kauppila_Harris_v2.docx"
)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------


def set_cell_border(cell, **kwargs):
    """Apply borders to a table cell (top/bottom/left/right)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = qn(f"w:{edge}")
            element = tcBorders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:color"), "BFBFBF")


def shade_cell(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def configure_document(doc: Document) -> None:
    """Set a clean, readable baseline: letter size, 1\" margins, Calibri 11."""
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for name, size, bold, color in [
        ("Title", 24, True, RGBColor(0x1F, 0x1F, 0x1F)),
        ("Heading 1", 16, True, RGBColor(0x1F, 0x1F, 0x1F)),
        ("Heading 2", 13, True, RGBColor(0x1F, 0x1F, 0x1F)),
        ("Heading 3", 11, True, RGBColor(0x40, 0x40, 0x40)),
    ]:
        if name in doc.styles:
            s = doc.styles[name]
            s.font.name = "Calibri"
            s.font.size = Pt(size)
            s.font.bold = bold
            s.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_para(
    doc: Document,
    text: str,
    *,
    italic: bool = False,
    bold: bool = False,
) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------


def add_budget_table(doc: Document, rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(1.8), Inches(4.2), Inches(0.8)]
    hdr_cells = table.rows[0].cells
    headers = ["Category", "Item", "Cost"]
    for i, (cell, header) in enumerate(zip(hdr_cells, headers)):
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        shade_cell(cell, "E8E8E8")
        cell.width = widths[i]
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)

    last_cat = None
    for cat, item, cost in rows:
        row = table.add_row().cells
        row[0].text = "" if cat == last_cat else cat
        row[1].text = item
        row[2].text = cost
        for i, cell in enumerate(row):
            cell.width = widths[i]
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
        if cost.upper() == "TOTAL" or item.upper() == "TOTAL":
            for cell in row:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.bold = True
                shade_cell(cell, "F3F3F3")
        last_cat = cat


def add_plan_table(doc: Document, rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [Inches(1.4), Inches(1.1), Inches(3.5), Inches(0.8)]
    headers = ["Phase", "Timeline", "Milestones", "Cost"]
    hdr = table.rows[0].cells
    for i, (cell, text) in enumerate(zip(hdr, headers)):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        shade_cell(cell, "E8E8E8")
        cell.width = widths[i]
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)

    for phase, timeline, milestones, cost in rows:
        row = table.add_row().cells
        row[0].text = phase
        row[1].text = timeline
        row[2].text = milestones
        row[3].text = cost
        for i, cell in enumerate(row):
            cell.width = widths[i]
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    # ---- Title block ---------------------------------------------------- #
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LACMA Art + Technology Lab")
    run.bold = True
    run.font.size = Pt(22)
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run("2026 Grant Application")
    run2.bold = True
    run2.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(18)
    meta_lines = [
        ("Applicants: ", True),
        ("Jesse Kauppila  &  Kameron Decker Harris", False),
        ("\nProject: ", True),
        ("Sunrise / Sunset", False),
        ("\nProject Website: ", True),
        ("sunrisesunset.studio", False),
        ("\nProject Repository: ", True),
        ("github.com/jessekauppila/the-sunset-webcam-map", False),
        ("\nApplication Deadline: ", True),
        ("April 22, 2026, 11:59 PM PST", False),
        ("\n\nJesse Kauppila: ", True),
        ("jessekauppila.art  ·  github.com/jessekauppila", False),
        ("\nKameron Decker Harris: ", True),
        ("glomerul.us/research.html  ·  Western Washington University", False),
    ]
    for text, bold in meta_lines:
        r = meta.add_run(text)
        r.bold = bold

    # ---- 1. Name of Your Project ---------------------------------------- #
    add_heading(doc, "1. Name of Your Project", level=1)
    add_para(
        doc,
        "Sunrise / Sunset: A Real-Time Live Stream of Sunrises and Sunsets "
        "as They Travel Around the World.",
    )

    # ---- 2. Three Words ------------------------------------------------- #
    add_heading(doc, "2. Three Words", level=1)
    add_para(doc, "Sunrise. Sunset. Forever.", bold=True)

    # ---- 3. One-Sentence Description ------------------------------------ #
    add_heading(doc, "3. One-Sentence Description", level=1)
    add_para(
        doc,
        "A perpetual stream of real, current sunrises and sunsets — pulled "
        "from existing webcam APIs and supplemented by custom edge-computing "
        "cameras, archived and ranked by a neural network learning to see "
        "beauty.",
    )

    # ---- 4. Full Description -------------------------------------------- #
    # Target: ~498 words (under 500). Trimmed from 655.
    add_heading(doc, "4. Full Description of the Proposed Project", level=1)

    add_para(
        doc,
        "In a cultural moment defined by AI-generated imagery and AI-waged "
        "warfare, we have become increasingly divorced from the beauty that "
        "already exists in the real world around us. This project inverts "
        "that relationship: it uses AI to find real sunrises and sunsets "
        "happening right now, somewhere on Earth.",
    )
    add_para(
        doc,
        "Sunrise / Sunset operates across four layers: a network of "
        "webcams, a web application, a neural-network model, and a gallery "
        "installation. Together, they locate webcams currently showing "
        "good sunrises and sunsets, archive them, and display them.",
    )
    add_para(
        doc,
        "The work is inspired by Christian Marclay’s The Clock and Janet "
        "Cardiff’s Forty-Part Motet — sublime wholes composed of many "
        "simultaneous parts. I hope to touch something beyond any single "
        "sunrise or sunset — the idea, corny as it sounds, that there is "
        "always beauty somewhere in the world.",
    )
    add_para(
        doc,
        "The web application queries webcam APIs and uses astronomical "
        "calculations (subsolar-point geometry) to locate cameras near the "
        "terminator — where day turns into night. It is the hub for the "
        "rest of the project: browsing webcams, rating them, and navigating "
        "the archive.",
    )
    add_para(
        doc,
        "I am also building small, custom edge-computing cameras from "
        "Raspberry Pi Zeros that provide better sunrise and sunset coverage "
        "than commercial APIs. They can be tuned remotely, streaming only "
        "when a sunset worth capturing is developing. In a second life, "
        "they could be offered as a Kickstarter object — funding the "
        "project while distributing its sensor network to participants.",
    )
    add_para(
        doc,
        "The machine-perception layer, led by Kameron, trains models on a "
        "continuous 0.0–1.0 aesthetic scale, encoding judgment about light, "
        "atmosphere, and landscape into a neural network. The pipeline "
        "combines my own ratings with assessments from vision-language "
        "models and Creative Commons sunset photography from the web. The "
        "AI’s scores shape the installation: higher-rated tiles render "
        "larger in the mosaic, so machine aesthetic judgment becomes "
        "visible as form.",
    )
    add_para(
        doc,
        "The current installation is a diptych displaying live sunrise "
        "and sunset mosaics. I’m working toward a "
        "ring of inward-facing monitors showing every current sunrise and "
        "sunset along the terminator — positioning the viewer at the "
        "center of the planet’s day/night boundary.",
    )
    add_para(
        doc,
        "With LACMA’s support, we propose to expand this work in four "
        "directions:",
    )
    add_bullets(
        doc,
        [
            "Scale the installation into the full ring.",
            "Expand the sensor network through partnerships with "
            "commercial, educational, and open-source webcam providers — "
            "and by deploying bespoke edge cameras.",
            "Deepen the machine-learning pipeline toward real-time, "
            "image-based inference.",
            "Build a public rating app — a “dating app for sunsets” — "
            "where anyone online rates them alongside the AI, making "
            "visible the many kinds of beauty in sunsets and the "
            "human–machine perception gap.",
        ],
    )
    add_para(
        doc,
        "The project sits at the intersection of landscape art, network "
        "culture, found photography and video, and AI. It connects to "
        "LACMA’s 1967 Art and Technology legacy of light and perception "
        "at planetary scale.",
    )
    add_para(
        doc,
        "I also think it would be wonderful to stand at the center of a "
        "ring of sunrises and sunsets — feeling the turning of the world, "
        "sunsets changing quickly at the equator and slowly at the poles. "
        "Who doesn’t love a sunset?",
    )

    # ---- 5. Bio --------------------------------------------------------- #
    add_heading(doc, "5. Bio", level=1)
    add_para(
        doc,
        "Jesse Kauppila is an artist and engineer based in Bellingham, WA. "
        "During an MFA at Carnegie Mellon University, I learned to program "
        "robotic arms for an art project, and I used that skill to begin a "
        "new career path — one that led to fabricating work for Anish "
        "Kapoor and Charles Ray, then to aerospace at Joby Aviation (the "
        "electric air-taxi company), and, more recently, to building AI "
        "data centers for big tech.",
    )
    add_para(
        doc,
        "I am now returning to my art practice, applying what I have "
        "learned in the tech world to projects like this one. Financial "
        "security has given me the room to clarify my values and focus on "
        "making work I genuinely believe in. For this project I am "
        "collaborating with Kameron Decker Harris, a computational "
        "scientist at Western Washington University, who leads the "
        "machine-learning architecture.",
    )
    add_para(
        doc,
        "Outside of studio and professional work, I ski, climb, bike, and "
        "run in the mountains of the Pacific Northwest. My relationship "
        "with nature informs my interest in using technology to cultivate "
        "a deeper connection to the world we all live in.",
    )
    add_para(
        doc,
        "Kameron Decker Harris is a computational scientist and applied "
        "mathematician at Western Washington University whose research "
        "focuses on networked dynamical systems — both biological and "
        "artificial neural networks. His work sits at the intersection of "
        "machine learning, dynamical systems, graph theory, and "
        "statistical inference. Prior collaborations with the Allen "
        "Institute for Brain Science advanced whole-brain connectivity "
        "inference from viral-tracing experiments, and he has published "
        "on sparsity-driven learning in biologically-inspired "
        "random-feature networks (including work showing V1-like tuning "
        "properties improving image recognition). For Sunrise / Sunset, "
        "Kameron leads the machine-learning architecture and training "
        "strategy, and collaborates on interpreting the human–machine "
        "aesthetic gap that the project makes explicit.",
    )

    # ---- 6. Artistic or Creative Merit ---------------------------------- #
    add_heading(doc, "6. Artistic or Creative Merit", level=1)
    add_para(
        doc,
        "Sunrise / Sunset points machine perception at actual sunrises and "
        "sunsets currently happening on Earth — not at machine-generated "
        "imaginings of them. It amplifies a universally shared aesthetic "
        "experience by combining dozens of simultaneous real views — drawn "
        "from public webcams and custom cameras of our own — into a single "
        "planetary portrait, curated by an AI trained to recognize which "
        "sunsets to highlight. The work proposes that the most interesting "
        "artistic use of AI may not be generation but attention: directing "
        "machine vision toward the sublime that already exists, "
        "continuously, in the real world.",
    )

    # ---- 7. Dialogue Between Technology and Culture --------------------- #
    add_heading(doc, "7. Dialogue Between Technology and Culture", level=1)
    add_para(
        doc,
        "We are in a cultural moment defined by AI-generated imagery and "
        "AI that wages war. These algorithms strip the world of its "
        "beauty, texture, and context. This project seeks to invert that "
        "relationship. Instead of asking a machine to imagine a sunset, "
        "it asks a machine to recognize the real ones happening right "
        "now, somewhere on Earth. It proposes that the most interesting "
        "artistic use of machine perception may not be generation but "
        "attention — directing the machine’s gaze toward the sublime "
        "that already exists, continuously, in the literal world. A tool "
        "for noticing rather than inventing.",
    )

    # ---- 8. Public Engagement Plan -------------------------------------- #
    # Target: ~100 words. Trimmed from 120.
    add_heading(doc, "8. Public Engagement Plan", level=1)
    add_para(
        doc,
        "The web application is live (moving to sunrisesunset.studio), "
        "offering global access to the terminator map and mosaics. The "
        "installation is portable and gallery-ready. Five public "
        "touchpoints:",
    )
    add_bullets(
        doc,
        [
            "The live web app and a rating site — a “dating app for "
            "sunsets” — where anyone online rates sunsets alongside the "
            "AI, feeding the model in real time.",
            "An early two-screen exhibit at Canopy Art & Iron (Bow, WA).",
            "A mid-term demo at the 2027 LACMA Biennial Symposium.",
            "An open-source release of the ML pipeline, model, and archive.",
            "A Kickstarter campaign for the custom edge cameras — funding "
            "plus distributed participation.",
        ],
    )

    # ---- 9. Other Funding Sources and In-Kind Support ------------------- #
    add_heading(doc, "9. Other Funding Sources and In-Kind Support", level=1)
    add_para(
        doc,
        "To date, Jesse and Kameron have jointly self-funded this project "
        "at approximately $1,500 — covering web hosting, installation and "
        "display hardware, early webcam prototyping, and AI/software "
        "subscriptions and API tokens.",
    )
    add_para(
        doc,
        "Potential partnerships to pursue during the grant period:",
        bold=True,
    )
    add_bullets(
        doc,
        [
            "Windy.com — commercial webcam API access; a partnership or "
            "sponsored data-access agreement is a stronger fit for this "
            "project than a paid subscription, and we would pursue one "
            "during the grant period.",
            "Other webcam networks and providers — potential partnerships "
            "with EarthCam, Skyline Webcams, AlpineWebcams, university "
            "and national-park live-feed programs, and open-source webcam "
            "projects, to expand the sensor network beyond any single "
            "commercial API.",
            "Snap Inc. and Anthropic (LACMA Lab partner companies) — the "
            "project’s use of LLM vision and real-time geospatial media "
            "overlaps with their technical interests; mentorship or "
            "in-kind API access are natural collaboration paths.",
        ],
    )
    add_para(doc, "Potential follow-up funding:", bold=True)
    add_bullets(
        doc,
        [
            "A Kickstarter campaign for the custom edge-computing cameras "
            "— simultaneously funding continued development and "
            "distributing the sensor network into the hands of backers.",
            "Possible retail sales of the cameras through LACMA’s gift "
            "shop or similar cultural-institution channels.",
        ],
    )
    add_para(doc, "No other funding is currently committed.")

    # ---- 10. Total Amount Requested ------------------------------------- #
    add_heading(doc, "10. Total Amount Requested", level=1)
    add_para(doc, "$50,000", bold=True)

    # ---- 11. Detailed Project Budget ------------------------------------ #
    add_heading(doc, "11. Detailed Project Budget", level=1)

    webcam_cat = (
        "Custom Webcam Prototyping (North America pilot — reference "
        "model for future global rollout)"
    )
    travel_cat = (
        "Travel (Bellingham, WA → Los Angeles — staying in LA for "
        "each event)"
    )

    budget_rows = [
        ("Principal Fees",
         "Jesse Kauppila — artist time: research, development, "
         "installation, documentation (24 months)", "$17,000"),
        (webcam_cat,
         "30 × Raspberry Pi Zero 2 W units (~$25 ea)", "$750"),
        (webcam_cat,
         "30 × Camera modules (wide-angle + standard) (~$20 ea)", "$600"),
        (webcam_cat,
         "30 × MicroSD cards, high-endurance 16–32 GB (~$12 ea)", "$360"),
        (webcam_cat,
         "30 × Power supplies and cables (~$10 ea)", "$300"),
        (webcam_cat,
         "30 × Weatherproof enclosures (~$15 ea)", "$450"),
        (webcam_cat,
         "Custom PCB design + small-run fabrication (Kickstarter-ready)",
         "$1,500"),
        (webcam_cat,
         "Enclosure design, 3D printing, prototyping iterations",
         "$1,000"),
        (webcam_cat,
         "Cellular / LTE modems for remote-deployment units (5–10 units)",
         "$500"),
        (webcam_cat,
         "Shipping to North American deployment partners, tools, "
         "spare parts", "$540"),
        (webcam_cat,
         "Webcam subtotal (30 units for NA pilot)", "$6,000"),
        ("Installation Hardware",
         "4–6 additional 27\" portrait monitors for ring installation",
         "$2,700"),
        ("Installation Hardware",
         "Additional Raspberry Pi 4B units, cases, SD cards, HDMI cables",
         "$800"),
        ("Installation Hardware",
         "Dedicated GPU workstation or equivalent cloud GPU credits "
         "(ML training)", "$2,500"),
        ("Software & Services",
         "Anthropic Claude subscription ($100/mo × 24 months)", "$2,400"),
        ("Software & Services",
         "Vercel Pro ($20/mo × 24 months)", "$480"),
        ("Software & Services",
         "Cursor Pro ($20/mo × 24 months)", "$480"),
        ("Software & Services",
         "Neon Postgres (24 months, estimated)", "$600"),
        ("Software & Services",
         "Mapbox GL JS — free tier covers expected usage (50,000 "
         "monthly map loads); overage, if any, from contingency",
         "$0"),
        ("Software & Services",
         "OpenAI + Google Gemini vision API calls (LLM labeling "
         "pipeline)", "$1,000"),
        ("Software & Services",
         "Domain registration + SSL (sunrisesunset.studio, 2 years)",
         "$100"),
        ("Installation Materials",
         "Monitor stands / mounts for portrait and ring configuration",
         "$1,500"),
        ("Installation Materials",
         "Wiring, cable management, power distribution for multi-screen "
         "array", "$500"),
        (travel_cat,
         "2027 Biennial Symposium trip — airfare (BLI/SEA ↔ LAX), 6 "
         "nights LA lodging, ground transport, meals, hardware shipping",
         "$2,700"),
        (travel_cat,
         "2028 Biennial Demo Day trip — airfare, 8 nights LA lodging, "
         "ground transport, meals, ring-installation hardware shipping",
         "$3,400"),
        (travel_cat,
         "Canopy Art & Iron exhibit (Bow, WA — local drive from "
         "Bellingham) — fuel, supplies, install day", "$300"),
        ("Co-applicant + Collaborator Fees",
         "Kameron Decker Harris — co-lead compensation: ML architecture, "
         "training strategy, model evaluation (24 months)", "$2,000"),
        ("Co-applicant + Collaborator Fees",
         "Installation / fabrication assistance", "$2,000"),
        ("Documentation",
         "Video + photo documentation of installation and process",
         "$1,500"),
        ("Contingency",
         "Unforeseen costs, replacement hardware, API overages (incl. "
         "any Mapbox / Firebase overages), shipping corrections",
         "$2,040"),
        ("TOTAL", "", "$50,000"),
    ]
    add_budget_table(doc, budget_rows)

    add_para(
        doc,
        "If any subscription lines come in lower than expected (for "
        "example, Anthropic offers credits through the LACMA "
        "partnership), the surplus flows into principal fees or "
        "additional webcam units for the pilot.",
        italic=True,
    )

    # ---- 12. Implementation Plan --------------------------------------- #
    add_heading(
        doc,
        "12. Implementation Plan — Milestones, Dates, and Costs",
        level=1,
    )

    plan_rows = [
        (
            "Phase 1: Foundation + Webcam v1",
            "Fall 2026 (Months 1–4)",
            "Formalize the co-lead agreement with Kameron Decker Harris "
            "and kick off ML architecture and training-strategy work. "
            "Upgrade ML pipeline to image-based production inference. "
            "Expand training dataset. Validate model performance against "
            "human ratings (Pearson > 0.80 gate). Initiate outreach to "
            "Windy, EarthCam, and other webcam providers regarding "
            "partnership / data access. Design and build first batch of "
            "5–10 custom Pi Zero 2 W cameras and identify North American "
            "deployment partners (universities, parks, independent "
            "camera hosts).",
            "$11,000",
        ),
        (
            "Phase 2: Ring Development + NA Webcam Rollout",
            "Winter 2026–27 (Months 4–8)",
            "Design and begin building the ring-monitor installation "
            "(6–8 portrait monitors). Develop the public rating app "
            "(“dating app for sunsets”). Iterate webcam hardware: custom "
            "PCB, weatherproof enclosure, remote-tuning firmware. Deploy "
            "remaining pilot webcam units across 8–10 North American "
            "locations (coastal, mountain, desert, Pacific Northwest); "
            "this pilot serves as the reference model for a later "
            "worldwide rollout. Exhibit the existing two-screen diptych "
            "— the sunrise / sunset pair — at Canopy Art & Iron’s "
            "annual rendezvous in Bow, WA. The diptych is more portable "
            "than the ring and is already the current working form of "
            "the installation, making Canopy’s annual event a natural, "
            "local venue for an early public showing before the LACMA "
            "symposium. Canopy’s “rooted in a sense of place” ethos "
            "complements the geographic, terminator-driven nature of "
            "the work. Gather audience response and document for LACMA.",
            "$13,000",
        ),
        (
            "Phase 3: LACMA Symposium Demo",
            "Spring 2027 (Months 8–12)",
            "Install and demonstrate at the LACMA 2027 Biennial "
            "Symposium. Present mid-term findings: model evolution, "
            "human vs. machine rating divergence, pilot-webcam data "
            "from deployed North American nodes, and reflections from "
            "the Canopy exhibit (and any other showings). Travel from "
            "Bellingham to LA.",
            "$9,000",
        ),
        (
            "Phase 4: Refinement + Kickstarter Prep + Global Network",
            "Summer–Fall 2027 (Months 12–18)",
            "Incorporate symposium feedback. Train a second-generation "
            "model on an expanded dataset (gallery-visitor ratings + "
            "rating-app data + pilot-webcam imagery). Finalize "
            "Kickstarter-ready camera design and campaign materials, "
            "using the NA pilot as a proven reference model. Use "
            "Kickstarter to fund and distribute additional cameras "
            "internationally. Open-source release of the ML pipeline.",
            "$8,000",
        ),
        (
            "Phase 5: Final Installation + 2028 Demo Day",
            "Winter–Spring 2028 (Months 18–24)",
            "Full-scale ring-monitor installation for 2028 Demo Day. "
            "Final documentation. Published dataset, model weights, and "
            "camera hardware plans. Travel from Bellingham to LA.",
            "$9,000",
        ),
    ]
    add_plan_table(doc, plan_rows)

    # ---- Supporting Images ---------------------------------------------- #
    add_heading(doc, "Supporting Images", level=1)
    add_para(
        doc,
        "To be embedded as JPEGs in the final submission document:",
        italic=True,
    )
    add_bullets(
        doc,
        [
            "Screenshot of the live map / globe view showing the "
            "terminator and webcam markers.",
            "Screenshot of the sunrise and sunset mosaics side by side "
            "(the current two-screen diptych).",
            "Rendering or mock-up of the ring-monitor installation with "
            "a viewer at the center.",
            "Photograph of an early Raspberry Pi Zero webcam prototype.",
            "Diagram of the ML pipeline architecture, or an example "
            "mosaic showing tile-size variation driven by AI quality "
            "scores.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
